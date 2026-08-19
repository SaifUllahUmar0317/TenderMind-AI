import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def clean_xml_string(s: str) -> str:
    """Removes control characters that are invalid in XML 1.0."""
    if not s:
        return ""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\x84\x86-\x9F]', '', s)

def set_cell_background(cell, hex_color: str):
    """Sets background shading color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding margins in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, hex_color="CBD5E1"):
    """Adds subtle clean borders to a Word table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{hex_color}"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{hex_color}"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

class DocumentExporter:
    """
    Document Exporter for generating clean, highly formatted .txt and .docx files
    without report wrapper metadata.
    """

    @staticmethod
    def generate_txt(extraction_result: dict, output_path: str) -> str:
        """Writes extracted text content into a clean text file."""
        lines = []
        for page in extraction_result.get("pages", []):
            lines.append(clean_xml_string(page.get("text", "")).strip())
            lines.append("\n")

        full_content = "\n".join(lines).strip()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_content)

        return output_path

    @classmethod
    def generate_docx(cls, extraction_result: dict, output_path: str) -> str:
        """
        Generates a clean Word (.docx) document containing ONLY the original document content
        formatted with native Word headings, lists, key-value pairs, and tables.
        """
        doc = Document()

        # Page Setup - Standard Margins
        for section in doc.sections:
            section.top_margin = Inches(0.85)
            section.bottom_margin = Inches(0.85)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

        # Palette Colors
        COLOR_PRIMARY = RGBColor(15, 23, 42)      # Dark Slate #0F172A
        COLOR_ACCENT = RGBColor(79, 70, 229)     # Royal Indigo #4F46E5
        COLOR_BODY = RGBColor(51, 65, 85)        # Body Charcoal #334155

        pages = extraction_result.get("pages", [])

        for idx, page in enumerate(pages):
            if idx > 0:
                doc.add_page_break()

            raw_text = clean_xml_string(page.get("text", ""))
            lines = raw_text.split("\n")

            in_markdown_table = False
            table_lines_buffer = []

            for line in lines:
                line_str = line.rstrip()

                # Detect Markdown Table Row
                if line_str.startswith("| ") and " |" in line_str:
                    in_markdown_table = True
                    table_lines_buffer.append(line_str)
                    continue
                else:
                    if in_markdown_table:
                        cls._render_native_word_table(doc, table_lines_buffer)
                        table_lines_buffer = []
                        in_markdown_table = False

                if not line_str.strip():
                    continue

                line_trimmed = line_str.strip()

                # 1. Heading Detection
                # Markdown headers (# , ## , ###)
                if line_trimmed.startswith("#"):
                    level = min(len(line_trimmed) - len(line_trimmed.lstrip("#")), 3)
                    clean_h = line_trimmed.lstrip("#").strip()
                    cls._add_heading(doc, clean_h, level=level)

                # ALL CAPS lines (e.g., "AI INTEGRATION PROPOSAL", "TABLE OF CONTENTS")
                elif len(line_trimmed) <= 65 and line_trimmed.isupper() and len(line_trimmed) > 3 and not line_trimmed.startswith("HTTP"):
                    cls._add_heading(doc, line_trimmed, level=1)

                # Numbered / Section Heading lines (e.g., "1. Executive Summary", "SECTION 2: ...")
                elif re.match(r'^(?:SECTION\s+\d+|CHAPTER\s+\d+|\d+\.\d+|\d+\.)\s+[A-Z]', line_trimmed):
                    cls._add_heading(doc, line_trimmed, level=2)

                # 2. List Detection
                # Bullet Lists (•, ●, ▪, ◆, ➢, -, *, +, —, –)
                bullet_match = re.match(r'^(?:[•●▪◆➢\-\*\+—–]|o\b)\s*(.+)', line_trimmed)
                
                # Numbered Lists (1., 1.1, 1), (1), A., a), (a))
                num_match = re.match(r'^((?:\(?\d+(?:\.\d+)*[\.\)]|\(?[a-zA-Z][\.\)]|\(?\b[ivxXILMC]+\b[\.\)]))\s+(.+)', line_trimmed)

                if bullet_match:
                    content = bullet_match.group(1).strip()
                    cls._add_bullet_item(doc, content, COLOR_BODY)

                elif num_match:
                    num_prefix = num_match.group(1).strip()
                    content = num_match.group(2).strip()
                    cls._add_numbered_item(doc, num_prefix, content, COLOR_PRIMARY, COLOR_BODY)

                # 3. Key-Value Pairs ("Label: Value")
                elif ":" in line_trimmed and not line_trimmed.startswith("http"):
                    colon_idx = line_trimmed.find(":")
                    key_part = line_trimmed[:colon_idx+1]
                    val_part = line_trimmed[colon_idx+1:]

                    if len(key_part) <= 35:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(3)
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.15

                        r_key = p.add_run(key_part + " ")
                        r_key.font.name = "Calibri"
                        r_key.font.size = Pt(11)
                        r_key.font.bold = True
                        r_key.font.color.rgb = COLOR_PRIMARY

                        r_val = p.add_run(val_part.strip())
                        r_val.font.name = "Calibri"
                        r_val.font.size = Pt(11)
                        r_val.font.color.rgb = COLOR_BODY
                    else:
                        cls._add_standard_paragraph(doc, line_trimmed, COLOR_BODY)

                else:
                    # Standard Paragraph Text
                    cls._add_standard_paragraph(doc, line_trimmed, COLOR_BODY)

            # Flush remaining table buffer if page ended with table
            if in_markdown_table and table_lines_buffer:
                cls._render_native_word_table(doc, table_lines_buffer)

        doc.save(output_path)
        return output_path

    @staticmethod
    def _add_heading(doc, text: str, level: int = 1):
        """Adds a clean formatted heading with proper level, spacing, and styling."""
        COLOR_PRIMARY = RGBColor(15, 23, 42)    # Dark Slate
        COLOR_ACCENT = RGBColor(79, 70, 229)   # Royal Indigo
        COLOR_SUB = RGBColor(51, 65, 85)        # Body Charcoal

        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.bold = True

        if level == 1:
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            r.font.size = Pt(18)
            r.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r.font.size = Pt(14)
            r.font.color.rgb = COLOR_ACCENT
        else:
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            r.font.size = Pt(12.5)
            r.font.color.rgb = COLOR_SUB

    @staticmethod
    def _add_bullet_item(doc, content: str, color: RGBColor):
        """Adds a clean indented bullet list item in Word with hanging indent."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.20)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

        # Bullet symbol with tab stop for crisp alignment
        r_bullet = p.add_run("•\t")
        r_bullet.font.name = "Arial"
        r_bullet.font.size = Pt(11)
        r_bullet.font.bold = True
        r_bullet.font.color.rgb = RGBColor(79, 70, 229)  # Royal Indigo

        r_text = p.add_run(content)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = color

    @staticmethod
    def _add_numbered_item(doc, num_prefix: str, content: str, color_num: RGBColor, color_text: RGBColor):
        """Adds a clean indented numbered list item in Word with hanging indent."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.40)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15

        r_num = p.add_run(num_prefix + "\t")
        r_num.font.name = "Calibri"
        r_num.font.size = Pt(11)
        r_num.font.bold = True
        r_num.font.color.rgb = color_num

        r_text = p.add_run(content)
        r_text.font.name = "Calibri"
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = color_text

    @staticmethod
    def _add_standard_paragraph(doc, text: str, color: RGBColor):
        """Helper to add clean body paragraph."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = color

    @staticmethod
    def _render_native_word_table(doc, table_lines: list):
        """
        Parses Markdown table lines into a formatted Native Word Table.
        """
        parsed_rows = []
        for line in table_lines:
            if re.match(r'^\|(?:\s*:?-+:?\s*\|)+$', line.strip()):
                continue
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if any(cells):
                parsed_rows.append(cells)

        if not parsed_rows:
            return

        col_count = max(len(r) for r in parsed_rows)
        row_count = len(parsed_rows)

        word_table = doc.add_table(rows=row_count, cols=col_count)
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(word_table, "CBD5E1")

        for r_idx, row_cells in enumerate(parsed_rows):
            padded_cells = row_cells + [""] * (col_count - len(row_cells))

            for c_idx, cell_value in enumerate(padded_cells):
                cell = word_table.cell(r_idx, c_idx)
                set_cell_margins(cell, top=100, bottom=100, left=140, right=140)

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                r = p.add_run(cell_value)
                r.font.name = "Calibri"

                if r_idx == 0:
                    set_cell_background(cell, "1E293B")
                    r.font.size = Pt(10)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                    set_cell_background(cell, bg_color)
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(51, 65, 85)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)
